"""
Render a generated report dict to a `.docx`.

The flow is now HTML-first:
    LLM markdown  →  markdown_to_html()  →  htmldocx.HtmlToDocx() inside python-docx

python-docx handles the title page (custom run formatting, RGB colours)
and section headings. The body of each section comes from the section's
sanitized HTML via htmldocx, which preserves GFM tables, nested lists,
bold runs, hyperlinks, and right-aligned numeric columns — things the
previous hand-rolled markdown→DOCX converter flattened.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.shared import Pt, RGBColor
from htmldocx import HtmlToDocx

from .html_renderer import markdown_to_html


logger = logging.getLogger(__name__)


def _add_title_page(doc: Document, report: Dict[str, Any]) -> None:
    """Cover header with entity name + FY range + audience line."""
    title = doc.add_heading("Credit Analysis Report", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    entity = doc.add_paragraph()
    entity_run = entity.add_run(report.get("entity_name", "—"))
    entity_run.font.size = Pt(16)
    entity_run.bold = True

    fys = report.get("fys") or []
    if fys:
        p = doc.add_paragraph()
        p.add_run(f"Financial years reviewed: {', '.join(fys)}").italic = True

    p = doc.add_paragraph()
    p.add_run("Prepared for: Senior Management / Credit Approval Committee").italic = True

    generated_at = (report.get("generated_at") or "")[:10]
    if generated_at:
        p = doc.add_paragraph()
        p.add_run(f"Generated: {generated_at}").italic = True


def _add_section(doc: Document, parser: HtmlToDocx, section: Dict[str, Any]) -> None:
    """Add one section's heading + HTML body to the document."""
    number = section.get("number")
    title = section.get("title", "")
    heading_text = f"{number}. {title}" if number is not None else title
    doc.add_heading(heading_text, level=1)

    html = section.get("html") or markdown_to_html(section.get("markdown", ""))
    if not html.strip():
        return
    try:
        parser.add_html_to_document(html, doc)
    except Exception:
        # Fail-soft: if htmldocx chokes on a malformed fragment, fall back
        # to the raw markdown as plaintext so the section isn't lost.
        logger.exception("htmldocx failed for section %s; falling back to plaintext", section.get("code"))
        doc.add_paragraph(section.get("markdown", "") or "")


def write_report_docx(report: Dict[str, Any], out_path: str | Path) -> Path:
    """Persist the report dict as a Word document."""
    doc = Document()
    _add_title_page(doc, report)
    doc.add_page_break()

    parser = HtmlToDocx()
    # Reasonable default table style for credit-memo content. Falls back
    # to "Table Grid" or to no style if the chosen one isn't on the
    # default Word.dotx (older Word installs).
    try:
        parser.table_style = "Light Grid Accent 1"
    except Exception:
        try:
            parser.table_style = "Table Grid"
        except Exception:
            pass

    for section in report.get("sections", []) or []:
        _add_section(doc, parser, section)
        doc.add_paragraph()  # blank-line spacer between sections

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
